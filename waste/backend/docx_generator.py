
Action: file_editor create /app/backend/docx_generator.py --file-text "\"\"\"DOCX Generator - creates ATS-compliant DOCX resumes\"\"\"
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(content: dict) -> io.BytesIO:
    \"\"\"Generate an ATS-friendly DOCX from optimized resume content\"\"\"
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Heading style
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.color.rgb = RGBColor(26, 26, 46)
        heading_style.font.name = 'Calibri'
    
    # Summary
    summary = content.get('summary', '')
    if summary:
        doc.add_heading('PROFESSIONAL SUMMARY', level=2)
        doc.add_paragraph(summary)
    
    # Skills
    skills = content.get('skills', [])
    if skills:
        doc.add_heading('SKILLS', level=2)
        skills_text = \" | \".join(skills) if isinstance(skills, list) else str(skills)
        doc.add_paragraph(skills_text)
    
    # Experience
    experience = content.get('experience', [])
    if experience:
        doc.add_heading('EXPERIENCE', level=2)
        for exp in experience:
            if isinstance(exp, dict):
                title = exp.get('title', '')
                company = exp.get('company', '')
                period = exp.get('period', '')
                
                # Job title line
                header = title
                if company:
                    header += f\" — {company}\"
                if period:
                    header += f\" ({period})\"
                
                p = doc.add_paragraph()
                run = p.add_run(header)
                run.bold = True
                run.font.size = Pt(11)
                
                # Bullet points
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet:
                        bp = doc.add_paragraph(bullet, style='List Bullet')
                        bp.paragraph_format.space_after = Pt(2)
    
    # Education
    education = content.get('education', [])
    if education:
        doc.add_heading('EDUCATION', level=2)
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                year = edu.get('year', '')
                
                line = degree
                if institution:
                    line += f\" — {institution}\"
                if year:
                    line += f\" ({year})\"
                
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(11)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
"
Observation: Create successful: /app/backend/docx_generator.py