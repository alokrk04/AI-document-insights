
import io
from PyPDF2 import PdfReader
from docx import Document


def parse_resume(content: bytes, file_type: str) -> str:
    \"\"\"Parse resume content from bytes based on file type\"\"\"
    if file_type == 'pdf':
        return parse_pdf(content)
    elif file_type == 'docx':
        return parse_docx(content)
    else:
        raise ValueError(f\"Unsupported file type: {file_type}\")


def parse_pdf(content: bytes) -> str:
    \"\"\"Extract text from PDF bytes\"\"\"
    reader = PdfReader(io.BytesIO(content))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return \"\n\".join(text_parts)


def parse_docx(content: bytes) -> str:
    \"\"\"Extract text from DOCX bytes\"\"\"
    doc = Document(io.BytesIO(content))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return \"\n\".join(text_parts)
"
