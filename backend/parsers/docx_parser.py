"""Parser for DOCX files using python-docx."""
from pathlib import Path


def parse_docx(file_path: Path) -> dict:
    """Extract text from DOCX preserving headings and structure."""
    from docx import Document

    doc = Document(str(file_path))
    sections = []
    full_text = ""
    current_section = "Document Start"

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            current_section = para.text.strip()
        if para.text.strip():
            sections.append({
                "page_number": 1,
                "section": current_section,
                "text": para.text.strip(),
                "source_type": "docx",
            })
            full_text += para.text.strip() + "\n\n"

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                sections.append({
                    "page_number": 1,
                    "section": current_section,
                    "text": row_text.strip(),
                    "source_type": "docx_table",
                })
                full_text += row_text.strip() + "\n"

    return {
        "pages": sections,
        "full_text": full_text.strip(),
        "page_count": 1,
        "source_type": "docx",
    }
